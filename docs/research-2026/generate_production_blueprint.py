#!/usr/bin/env python3
"""Generate PDF and DOCX for the expanded Autonomous Agentic Trading Production Blueprint."""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = SCRIPT_DIR
os.makedirs(OUT_DIR, exist_ok=True)

def sanitize(text):
    if not isinstance(text, str): return text
    return (text
        .replace("\u2014", "-").replace("\u2013", "-")
        .replace("\u2018", "'").replace("\u2019", "'")
        .replace("\u201c", '"').replace("\u201d", '"')
        .replace("\u2022", "-").replace("\u00a0", " "))

TITLE = "Autonomous Agentic Trading Production Blueprint 2026 (Expanded)"
DATE = "2026-06-14"
SUBTITLE = "For stocks, crypto, and all markets. True autonomy (hierarchical multi-agent, debate, reflection, memory, guardrails) - not bots or simple agents. From scratch to production."

SECTIONS = [
    ("Executive Summary", [
        "True autonomous agentic = multi-agent teams that perceive, debate/plan, act (real tradable orders with impact), reflect (regret/lessons), meta-learn (rule/strategy adaptation), and persist knowledge with minimal human intervention.",
        "Production requirements: Deterministic guardrails/kill-switches OUTSIDE LLM, realistic LOB/slippage/latency simulation, full audit (COT + episodes), multi-market data/execution, cost/latency control, reliability (resume, circuit breakers), regulatory compliance.",
        "2026 highlights: TradingAgents (LangGraph debate + memory + reflection + stock/crypto), QuantReplay (open multi-asset LOB matching engine + synthetic flow), Alpaca official MCP Server (natural language agent trading stocks/options/crypto), Nautilus Trader (Rust production event-driven), guardrails literature stressing runtime deterministic controls + kill switches.",
        "tredo/TREDO fit: Already has many production strengths (Rust DisciplinedCore, hierarchical groups + emerging debate, rich episodes + regret + meta, temporal loops, skills, trained memory recall, COT). Gaps: realistic sim, live broker (Alpaca MCP ideal), full multi-asset, advanced LOB execution modeling."
    ]),
    ("Core Architecture Layers for Autonomy", [
        "1. Perception (fast): Multi-source real-time (WS + broker feeds + Polygon + CCXT) + batch (fundamentals, on-chain, news, macro). Unified instrument model + regime/session awareness.",
        "2. Specialists (parallel): Technical, Fundamentals, Sentiment/Social, News/Macro, On-Chain, Options surface.",
        "3. Debate/Reasoning (medium): Structured states (bull/bear or Proposer/Critic/Risk/Historian). Aggregator with confidence/veto. Inject memory + verified context every step. Use LangGraph StateGraph (rapid) or custom Rust orchestrator.",
        "4. Decision/Execution: Synthesize -> Core rules validation -> Portfolio (cross-asset heat/correlations) -> Coordinator (order construction, impact estimate).",
        "5. Guardian/Safety (always): Hard rules + kill switch (global flag + mutex) + circuit breakers. Pre-action deterministic gate.",
        "6. Reflection/Meta (slow): Structured episode -> LLM reflection (regret 0-1, lessons, suggested changes) -> MetaControl proposes rule/strategy updates -> procedural memory.",
        "7. Memory tiers: Episodic (rich TradingEpisode), Vector (similar past recall), Procedural (lessons), Long-term external.",
        "8. Simulation/Validation: Essential. Use QuantReplay (multi-asset LOB + matching + synthetic flow + latency + slippage) or ABIDES-MARL/PyMarketSim for realism. Multi-stage: unit -> historical w/ costs -> walk-forward -> paper/sim -> limited live.",
        "9. Ops: Docker, OTEL + rich COT, resume from checkpoints, cost control (selective LLM), multi-level monitoring + alerts."
    ]),
    ("Multi-Market Data & Execution", [
        "Unified model: ticker + asset_class + suffix + session + microstructure.",
        "Feeds: Real-time (Alpaca/IBKR WS, Binance, Polygon depth), historical/fundamentals (normalized yf + paid), crypto on-chain + funding, news/sentiment (multi-source + summarize), macro/calendars, options chains/surfaces.",
        "Brokers: Alpaca (top for agents - official MCP Server 2026 with 61 endpoints for natural language + structured tools across stocks/options/crypto; high uptime; paper+live). IBKR (most complete global, 150+ order types, crypto in regions). CCXT (crypto unification).",
        "Execution realism: Model own order walking the LOB, variable slippage by liquidity/time, latency, partials, fees, borrow/funding. Paper always default with hard flags.",
        "MCP bridge: Expose tools so Claude/Cursor/custom agents can drive high-level strategy in English while core enforces rules."
    ]),
    ("Safety, Guardrails & Regulation (Non-Negotiable for Autonomy)", [
        "Probabilistic (LLM/debate/memory) proposes. Deterministic code disposes/validates/sizes/blocks.",
        "Mandatory: Action allowlists, pre-execution full rule checks (risk, confluence, session, correlation, exposure), step/loop limits + timeouts, global + per-strategy kill switch + mutex (instant halt + forensics), circuit breakers (data gaps, extremes, cost spikes), human gates for large/regret/rule changes.",
        "Regulatory: Deployers liable (no 'AI did it' defense). EU AI Act high-risk likely (transparency, oversight, logging). SEC/CFTC evolving - need explainability and controls. Design immutable full audit from day 1. Crypto adds custody/AML/on-chain issues.",
        "Production: Paper-only until long clean sim campaigns with regret analysis. Staged rollout (universe/size). Emergent behavior monitoring. Incident runbooks for agents."
    ]),
    ("Phased Roadmap (Scratch or Extend tredo)", [
        "Phase 0 (0-4w): Multi-asset ingestion + normalization, expanded DisciplinedCore, paper execution + accounting, basic episodes + reflection.",
        "Phase 1 (4-10w): Specialists + tools (parallel), full debate (states/aggregator), memory tiers + recall injection, temporal orchestrator, basic harness.",
        "Phase 2 (10-16w): Integrate realistic sim (QuantReplay or equiv for LOB/slippage/multi-asset), full data breadth (on-chain/options/FX/macro), hardened guardrails/kill/circuits, reflection+meta loop, reproducibility controls.",
        "Phase 3 (16-24w+): Live broker (Alpaca MCP first, then IBKR/CCXT), full OTEL + COT dashboards, cost/latency opt, deployment/monitoring/resume, regulatory audit package, tiny-capital live with gates.",
        "Phase 4: Cross-asset + portfolio agents, long-term memory scaling, MCP natural-language interface, continuous red-team + evolution.",
        "Gates: Every phase = paper + realistic sim validation + regret review. Live gate = 100+ clean cycles + positive expectancy after costs + full trails."
    ]),
    ("Tech Choices & tredo Alignment", [
        "Core: Rust (tredo/Nautilus style) for determinism, safety, perf, rules, loops, TUI, execution layer.",
        "Orchestration: LangGraph for rapid debate/state machine iteration (see TradingAgents patterns: parallel analysts, Invest/Risk debate states, memory log, reflection, checkpoints, asset_type stock/crypto). Port winners to Rust.",
        "Memory: redb + LanceDB (vector) + SQLite + external long-term.",
        "Sim: QuantReplay (or ABIDES-MARL/PyMarketSim) for production realism (LOB, synthetic flow, multi-asset, latency, slippage).",
        "LLM: Multi-provider (local Ollama + frontier), selective use, quick/deep models.",
        "tredo advantages: Already embodies many 2026 production recommendations (rules in code with memory adjustment, episodes+regret+meta, hierarchy+debate, temporal, COT, hybrid). Close sim/execution/multi-asset/MCP/guardrail-depth gaps and you have a differentiated safe autonomous reference system.",
        "Avoid for true production autonomy: Pure Python without strong deterministic guardrails layer; over-trusting historical backtests without LOB realism; weak kill/resume/audit."
    ]),
    ("Critical Risks & Mitigations (2026 Lessons)", [
        "Overfitting/adaptive search: Sealed test sets, report search budget, grounded deterministic data.",
        "Reproducibility: Pin dates, log full state, deterministic snapshots.",
        "Realism gap: Static slippage fails live. Use proper LOB simulators (QuantReplay etc.).",
        "Cost/latency: Debate multiplies LLM calls. Selective + cache + cheaper models.",
        "Safety escape: LLM bad proposal reaches market. Kill switch + pre-execution deterministic MUST exist.",
        "Regulatory: Opaque decisions without trails = liability. Audit everything.",
        "Ops: Feed/LLM/service outages. Graceful degradation + health + resume.",
        "From real attempts (TradingAgents docs, production posts, sim papers): Resolve instrument identity explicitly, inject memory reflections, checkpoint resume, separate asset pipelines, use synthetic flow for stress."
    ])
]

def make_pdf():
    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(100)
            self.cell(0, 8, sanitize("Autonomous Agentic Trading Production Blueprint 2026 | TREDO/tredo"), align="C", new_x="LMARGIN", new_y="NEXT")
            self.ln(1)
        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128)
            self.cell(0, 10, f"Page {self.page_no()}", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=14)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(20, 50, 100)
    pdf.multi_cell(0, 9, sanitize(TITLE), align="C")
    pdf.ln(1)

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(60)
    pdf.cell(0, 5, f"Date: {DATE}", new_x="LMARGIN", new_y="NEXT")
    pdf.multi_cell(0, 5, sanitize(SUBTITLE))
    pdf.ln(3)

    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(30)
    pdf.multi_cell(0, 4.5, sanitize("Comprehensive expanded research for building true autonomous agentic (hierarchical multi-agent, debate-driven, memory-learning, guardrail-enforced) trading systems from scratch or extending existing foundations like tredo. Covers stocks, crypto, FX, futures, options, multi-asset. Emphasis on production realism (simulation, execution, safety, ops) and autonomy (self-improvement loops) rather than simple bots."))
    pdf.ln(2)

    for h, bullets in SECTIONS:
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(20, 50, 100)
        pdf.multi_cell(0, 6, sanitize(h))
        pdf.ln(0.5)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(30)
        for b in bullets:
            pdf.set_x(pdf.l_margin + 2)
            pdf.multi_cell(0, 4.2, f"- {sanitize(b)}")
        pdf.ln(1.5)

    pdf.set_font("Helvetica", "I", 7)
    pdf.set_text_color(100)
    pdf.multi_cell(0, 4, sanitize("Key refs: arXiv Agentic Trading 2605.19337, TradingAgents (TauricResearch full graph + debate/memory/reflection), QuantReplay (multi-asset LOB sim), ABIDES-MARL/PyMarketSim/FinRL-Meta, Alpaca MCP Server 2026, Nautilus Trader (Rust), guardrails/kill-switch literature. Internal tredo docs/code (episodes, DisciplinedCore, hierarchy, temporal, skills). Not financial advice."))
    
    out = os.path.join(OUT_DIR, "agentic-trading-autonomous-production-blueprint-2026.pdf")
    pdf.output(out)
    print(f"PDF: {out}")
    return out

def make_docx():
    doc = Document()
    title = doc.add_heading(sanitize(TITLE), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"Date: {DATE}\n").italic = True
    p.add_run(sanitize(SUBTITLE)).italic = True

    doc.add_paragraph(sanitize("Expanded production blueprint for full autonomous agentic trading systems (stocks + crypto + all markets). Hierarchical multi-agent with debate, rich memory + reflection + meta-learning, deterministic guardrails outside LLM, realistic simulation (LOB/slippage), multi-asset execution, from-scratch-to-production roadmap. Builds on and references tredo/TREDO strengths while identifying precise gaps."))

    for h, bullets in SECTIONS:
        doc.add_heading(sanitize(h), level=1)
        for b in bullets:
            doc.add_paragraph(sanitize(b), style="List Bullet")

    doc.add_heading("Key References & Sources", level=1)
    refs = [
        "arXiv 2605.19337 - Agentic Trading survey (May 2026)",
        "TauricResearch/TradingAgents (full code: graph with debate states, memory log, reflection, tool nodes, checkpoints, asset_type stock/crypto, MCP relevance)",
        "QuantReplay (Quod Financial open-source multi-asset LOB matching + synthetic flow + latency/slippage)",
        "ABIDES-MARL, PyMarketSim, FinRL-Meta (realistic LOB/multi-agent/RL environments & benchmarks)",
        "Alpaca MCP Server (official 2026 - natural language + 61 structured endpoints for stocks/options/crypto agent trading)",
        "Nautilus Trader (Rust-native production event-driven multi-asset engine)",
        "Guardrails literature (runtime deterministic controls, kill switches, circuit breakers for agentic production)",
        "tredo/TREDO internal: AGENTIC_ARCHITECTURE_V2.md, DISCIPLINED_CORE.md, episodes/reflection/meta, skills, temporal loops, current paper/backtest stubs"
    ]
    for r in refs:
        doc.add_paragraph(sanitize(r), style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("\nNot financial, investment, or trading advice. Research artifact to support safe building of autonomous systems.")
    run.italic = True

    out = os.path.join(OUT_DIR, "agentic-trading-autonomous-production-blueprint-2026.docx")
    doc.save(out)
    print(f"DOCX: {out}")
    return out

if __name__ == "__main__":
    pdf = make_pdf()
    docx = make_docx()
    print(f"\nProduction blueprint reports generated in {OUT_DIR}")
    print("  - agentic-trading-autonomous-production-blueprint-2026.md (detailed source)")
    print(f"  - {os.path.basename(pdf) if pdf else 'PDF failed'}")
    print(f"  - {os.path.basename(docx) if docx else 'DOCX failed'}")
