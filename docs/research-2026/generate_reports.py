#!/usr/bin/env python3
"""
Generate PDF (fpdf2) and DOCX (python-docx). Sanitized for core fonts.
"""
import os
from datetime import datetime

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = SCRIPT_DIR if os.path.basename(SCRIPT_DIR) == "research" else os.path.join(os.getcwd(), "research")
os.makedirs(OUT_DIR, exist_ok=True)

def sanitize(text):
    if not isinstance(text, str):
        return text
    return (text
        .replace("\u2014", "-")   # em dash
        .replace("\u2013", "-")   # en dash
        .replace("\u2018", "'").replace("\u2019", "'")
        .replace("\u201c", '"').replace("\u201d", '"')
        .replace("\u2022", "-")
        .replace("\u00a0", " ")
    )

TITLE = sanitize("Agentic Trading Research Report - 2026")
DATE = "2026-06-14"
SUBTITLE = sanitize("Context: TREDO / tredo (Rust hierarchical multi-agent autonomous trading co-pilot)")

SECTIONS = [
    {
        "heading": sanitize("Executive Summary"),
        "bullets": [
            sanitize("Agentic trading: LLM-powered agents that perceive, reason/debate, act (trade), reflect, and adapt autonomously or semi-autonomously."),
            sanitize("2025-2026 explosion: open-source frameworks (TradingAgents, QuantAgent, AgenticTrading Lab) + mainstream products (Robinhood Agentic Trading)."),
            sanitize("Core 2026 themes: multi-agent debate simulating trading firms; guardrails (rules-as-code) outside LLM mandatory; episodic memory + reflection + meta-learning; hybrid stacks (orchestration + deterministic subs); reproducibility & regulatory challenges."),
            sanitize("tredo/TREDO alignment: Already implements many best practices (Rules + Memory > Prompting, two-tier agents + debate, rich TradingEpisodes + regret + trained memory recall, temporal loops, DisciplinedCore guardrails in Rust, full COT observability). Serves as strong reference impl for safe production agentic systems."),
        ],
        "body": sanitize("Agentic AI market ~$27.85B (2026). Finance adoption high but production low (11%). Papers and X discussions stress that without hard guardrails, agents lose money or create systemic risks. tredo's Rust-first disciplined approach is ahead of most pure-Python research prototypes.")
    },
    {
        "heading": sanitize("Key Frameworks & Examples"),
        "bullets": [
            sanitize("TradingAgents (TauricResearch, arXiv 2412.20138, v0.2.5 2026): LangGraph-based. Analysts (fundamentals/sentiment/news/technical) - Bull/Bear researchers (debate) - Trader - Risk/Portfolio Manager. Multi-LLM (incl. Grok, Ollama local). Persistent decision log + reflections + LangGraph checkpoints. Strong CLI/Docker/Python API."),
            sanitize("QuantAgent: Multi-agent (4 specialized) for HFT analysis from Stony Brook/CMU/Yale/UBC/Fudan. Parallel dimension analysis - synthesized trade decision (entry/exit/SL). Open-sourced 2026, significant X attention."),
            sanitize("Others: Open-Finance-Lab/AgenticTrading (educational platform); multiple arXiv (TradingGroup self-reflection, FactorMAD debate for factors, RMATS recursive typed agents); surveys on DRL-LLM-multi-agent evolution."),
            sanitize("Products: Robinhood Agentic Trading (2026) - third-party agents via MCP to dedicated accounts with safety controls. DeFi reports of agents managing substantial TVL in top pools."),
        ],
        "body": sanitize("These systems emphasize firm-like role specialization + debate. TradingAgents most directly comparable to tredo's architecture and actively maintained with broad LLM support.")
    },
    {
        "heading": sanitize("Architecture Patterns"),
        "bullets": [
            sanitize("Layered: Fast perception (prices, patterns) - Analysis specialists - Debate/reasoning (medium cadence) - Decision/execution - Guardian risk (always) - Slow reflection/meta (lessons, rule proposals)."),
            sanitize("Memory: Structured episodes (snapshot + trace + outcome + reflection with regret/lessons), vector similarity for 'similar past', long-term procedural (rule updates)."),
            sanitize("Orchestration: LangGraph dominant for stateful/auditable flows; custom (tredo) for safety-critical deterministic cores."),
            sanitize("Hybrid: Selective LLM (scarce), deterministic subs for speed/safety (pivots, risk calc), sidecar services for forecasts (cf. tredo Kronos). Full COT/decision audit trails required."),
        ],
        "body": sanitize("tredo's Fast/Med/Slow loops, two-tier (main+sub), Proposer/Critic/Risk/Historian debate, DisciplinedCore + apply_trained_memory, hierarchical recall (vector + agentmemory), and ratatui COT directly implement the patterns recommended across 2026 literature.")
    },
    {
        "heading": sanitize("Risks, Regulation & Challenges"),
        "bullets": [
            sanitize("Technical: Latency/cost (multiple LLM calls), non-determinism (sampling + live data), hallucinated facts without grounding, reproducibility hard (TradingAgents docs explicitly address this)."),
            sanitize("Risk: Emergent collusion, amplified flash events, overfitting via memory, operator confusion (paper vs live)."),
            sanitize("Regulatory/Legal: No-intent gap in manipulation law; deployer liability (negligence, agency, product); EU AI Act high-risk likely; need explainability + human oversight gates. Papers call for harm-based liability + private enforcement. X: stop-loss protocols for machine authority."),
            sanitize("Reproducibility note: Pin dates, ground data, use lower-temp non-reasoning models for repeatable backtests."),
        ],
        "body": sanitize("Sources repeatedly warn: pure agentic without unbreakable code guardrails will lose capital and create compliance nightmares. Audit everything (COT + episodes + logs).")
    },
    {
        "heading": sanitize("Recommendations & tredo Fit"),
        "bullets": [
            sanitize("Guardrails in deterministic code first (tredo DisciplinedCore excellent model)."),
            sanitize("Rich episodic memory + reflection + meta (tredo already has TradingEpisode + Reflector + MetaControl + trained recall)."),
            sanitize("Debate before action + specialist decomposition (tredo Phase C completing this)."),
            sanitize("Full observability (tredo ratatui TUI + COT is a standout)."),
            sanitize("Hybrid perf/safety lang (Rust core + Python glue - tredo's approach)."),
            sanitize("Paper-only validation + regret analysis before any real capital."),
            sanitize("For tredo: Finish debate+aggregator, upgrade vector to LanceDB, clean duplication, polish names/launcher/Docker, expand memory usage, more end-to-end validation loops."),
        ],
        "body": sanitize("tredo is not just another agentic trading project - it is a safety-first, memory-driven, observable realization of the exact architecture 2026 research advocates. It can serve as a reference for teams prioritizing production readiness over rapid Python prototyping.")
    },
]

def generate_pdf():
    try:
        from fpdf import FPDF
    except ImportError:
        print("fpdf2 not available")
        return None

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(100, 100, 100)
            self.cell(0, 8, sanitize("Agentic Trading Research 2026 | TREDO/tredo context"), align="C", new_x="LMARGIN", new_y="NEXT")
            self.ln(2)

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128)
            self.cell(0, 10, f"Page {self.page_no()}", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(20, 60, 120)
    pdf.multi_cell(0, 10, TITLE, align="C")
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(60)
    pdf.cell(0, 6, f"Date: {DATE}", new_x="LMARGIN", new_y="NEXT")
    pdf.multi_cell(0, 6, SUBTITLE)
    pdf.ln(4)

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(30)
    pdf.multi_cell(0, 5, sanitize("This report synthesizes 2026 web, academic, GitHub, and social sources on agentic trading systems (LLM/multi-agent autonomous trading). It highlights alignment with the tredo/TREDO codebase, which implements production-grade versions of recommended patterns (hierarchical agents, debate, guardrails in code, memory-driven self-improvement)."))
    pdf.ln(3)

    for sec in SECTIONS:
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_text_color(20, 60, 120)
        pdf.multi_cell(0, 7, sec["heading"])
        pdf.ln(1)

        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(30)
        pdf.multi_cell(0, 5, sec.get("body", ""))
        pdf.ln(1)

        pdf.set_font("Helvetica", "", 9)
        for b in sec["bullets"]:
            pdf.set_x(pdf.l_margin + 3)
            pdf.multi_cell(0, 5, f"- {b}")
        pdf.ln(2)

    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(100)
    pdf.multi_cell(0, 5, sanitize("Not financial advice. Research/educational use only. Sources include arXiv papers (TradingAgents 2412.20138, Agentic Trading 2605.19337), GitHub (TauricResearch/TradingAgents), X discussions, industry reports, Robinhood product pages."))
    pdf.ln(2)

    out = os.path.join(OUT_DIR, "agentic-trading-research-2026.pdf")
    pdf.output(out)
    print(f"PDF written: {out}")
    return out

def generate_docx():
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not available")
        return None

    doc = Document()

    title = doc.add_heading(TITLE, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"Date: {DATE}").italic = True
    p.add_run("\n" + SUBTITLE).italic = True

    doc.add_paragraph(sanitize("This report synthesizes 2026 sources on agentic trading (autonomous LLM/multi-agent systems for financial markets). Special attention to alignment with TREDO/tredo, a Rust-first hierarchical agentic trading co-pilot that already embodies many recommended 2026 best practices (guardrails in code, rich episodic memory + reflection, debate, temporal loops, full observability)."))

    for sec in SECTIONS:
        doc.add_heading(sec["heading"], level=1)
        if sec.get("body"):
            doc.add_paragraph(sec["body"])
        for b in sec["bullets"]:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Key References", level=1)
    refs = [
        "arXiv 2605.19337 (May 2026) - Agentic Trading: When LLM Agents Meet Financial Markets",
        "ResearchGate survey (June 2026) - Agentic Financial Trading Agents comprehensive review",
        "TradingAgents: arXiv 2412.20138; https://github.com/TauricResearch/TradingAgents (v0.2.5 2026)",
        "QuantAgent (2026 open-source multi-agent HFT from Stony Brook et al.)",
        "Robinhood Agentic Trading product (2026) - MCP-connected agents with dedicated accounts",
        "Internal: tredo Research.md, AGENTIC_ARCHITECTURE_V2.md, DISCIPLINED_CORE.md",
        "X posts & industry articles on frameworks, risks, and adoption (2026).",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("\nNot financial advice. For research and development of safe autonomous trading systems.")
    run.italic = True

    out = os.path.join(OUT_DIR, "agentic-trading-research-2026.docx")
    doc.save(out)
    print(f"DOCX written: {out}")
    return out

if __name__ == "__main__":
    pdf_path = generate_pdf()
    docx_path = generate_docx()
    print("\nDone. Files in:", OUT_DIR)
    print("  - agentic-trading-research-2026.md (source)")
    if pdf_path: print("  -", os.path.basename(pdf_path))
    if docx_path: print("  -", os.path.basename(docx_path))
