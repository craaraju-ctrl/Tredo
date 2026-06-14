#!/usr/bin/env python3
"""Generate PDF and DOCX for self-evolving agentic AI/trading research + intact system blueprint."""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = SCRIPT_DIR
os.makedirs(OUT_DIR, exist_ok=True)

def sanitize(text):
    if not isinstance(text, str): return text
    return (text
        .replace("\u2014", "-").replace("\u2013", "-")
        .replace("\u2018", "'").replace("\u2019", "'")
        .replace("\u201c", '"').replace("\u201d", '"')
        .replace("\u2022", "-").replace("\u00a0", " "))

TITLE = "Self-Evolving Agentic AI & Trading Systems - Intact System Blueprint 2026"
DATE = "2026-06-14"

SECTIONS = [
    ("Core Concept: Self-Evolving Agentic Systems", [
        "Self-evolving agents capture experience (successes, failures, regret), reflect (structured lessons), and update multiple layers: memory, tools, workflows, rules, even the improvement process itself.",
        "Key mechanisms: Reflection (Reflexion-style verbal RL), regret/outcome feedback, meta tool learning, intra-task replanning + inter-task batch adaptation, hierarchical memory as evolution substrate.",
        "2026 highlights: MetaAgent (self-reflection + autonomous tool construction without retraining), TradingGroup (per-agent self-reflection + dynamic risk in trading), HyperAgents (meta-agent improves the improver), broad surveys on what/when/how to evolve."
    ]),
    ("Self-Evolution in Agentic Trading (Specific)", [
        "TradingGroup (arXiv 2508.17565): Multi-agent (forecasting, style, decision) with explicit self-reflection that distills past successes/failures for analogous future scenarios + dynamic risk model for adaptive SL/TP/position sizing.",
        "TradingAgents: Decision log + realized returns/alpha → LLM-generated reflection → injection into future prompts (Portfolio Manager). Concrete example of closed experience → adaptation loop.",
        "Other patterns: Regret-driven rule updates, continual world models, evaluation agents feeding optimization suggestions back to signal/forecast agents, cross-asset reflective memory.",
        "Why powerful for trading: Clear delayed feedback (PnL, regret, alpha), recurring non-stationary patterns, high cost of repeating mistakes."
    ]),
    ("Diagnosis of Current Implementation Gaps", [
        "Strengths: Hierarchical groups + emerging debate with trained memory recall, rich TradingEpisode + PostTradeReflection (regret, lessons, suggested changes), DisciplinedCore (memory-adjustable rules in Rust), Reflector with deep reflection + recall, temporal loops, skills, COT.",
        "Gaps causing 'not working' for full self-evolving autonomy: Execution and backtester are stubs/placeholders; debate not fully wired end-to-end with robust aggregator; reflection loop not reliably feeding meta-control and procedural memory updates; no automatic realized-outcome + reflection injection on future runs (TradingAgents-style); limited realistic simulation for safe adaptation signals; multi-asset and dynamic risk adaptation incomplete.",
        "Result: Ambitious self-evolution vision exists in design and partial components, but the system cannot yet demonstrate closed-loop autonomous improvement."
    ]),
    ("The Intact System Target Blueprint (Complete Self-Evolving Design)", [
        "Layered architecture (build on tredo strengths): Perception (multi-asset verified feeds) → Specialists + Skills (deterministic fast) → Full Debate (Proposer/Critic/Risk/Historian with recall + aggregator) → Decision + Dynamic Risk (memory-informed adaptive SL/TP/sizing) → Execution (realistic paper then live via broker/MCP) → Outcome Capture (rich episodes) → Reflection (with prior memory context) → Memory Update (episodic + vector + procedural) → Meta-Control (high-regret patterns → rule/tool/workflow proposals with human gate on high impact) → Guardian/Safety (deterministic kill switches and checks outside evolving parts).",
        "Self-evolution loop (the 'intact' closed loop): Action + trace → realistic execution → structured reflection (regret + lessons) → memory update → meta layer proposes concrete adaptations → gated application → next cycle retrieves improved memory/rules → measurable improvement (lower regret on analogous situations, better risk-adjusted performance).",
        "Critical enablers: Realistic multi-asset LOB simulation (QuantReplay or equivalent) for safe evolution experiments; persistent decision log with automatic outcome resolution + reflection injection; versioned auditable rule/memory evolution; full COT + adaptation history observability.",
        "Multi-market: Asset-class aware (sessions, data densities, regimes) + cross-asset portfolio heat/correlations in risk and meta layers."
    ]),
    ("Practical Patterns, Safety & Roadmap", [
        "Good patterns: Ground reflections and proposals in retrieved past episodes; keep hard safety deterministic and non-evolvable; use both scalar regret and rich textual signals; version everything; evolve narrow then expand; combine intra-task replanning with inter-task meta updates.",
        "Safety: Deterministic guardrails (DisciplinedCore + kill switches + pre-action checks) must stay outside the self-evolving parts. Human gates on high-impact changes. Prevent adaptive overfitting with sealed evaluation and grounded data.",
        "Prioritized path to intact: 1. Complete debate + aggregator. 2. Realistic execution + paper modeling. 3. Closed outcome resolution + reflection injection loop. 4. Wire reflection → meta-control → rule/risk adaptation. 5. Integrate proper simulator. 6. Add system-level self-evolution metrics and monitoring. 7. Multi-asset breadth + live broker path (Alpaca MCP attractive).",
        "This gives a complete, working, highly adaptive autonomous agentic trading system that learns from mistakes and updates itself — the 'intact' target."
    ])
]

def make_pdf():
    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(100)
            self.cell(0, 7, sanitize("Self-Evolving Agentic AI & Trading - Intact System Blueprint 2026"), align="C", new_x="LMARGIN", new_y="NEXT")
            self.ln(1)
        def footer(self):
            self.set_y(-10)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(128)
            self.cell(0, 8, f"Page {self.page_no()}", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(20, 50, 100)
    pdf.multi_cell(0, 7, sanitize(TITLE), align="C")
    pdf.ln(1)

    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(60)
    pdf.cell(0, 5, f"Date: {DATE}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(30)
    intro = sanitize("Expanded research on self-evolving, reflective, meta-learning agentic AI and trading systems that learn from mistakes, update rules/memory/tools/workflows, and become highly adaptive over time. Includes diagnosis of current implementation gaps and a concrete 'intact system' blueprint for a complete working self-evolving autonomous trading architecture.")
    pdf.multi_cell(0, 4, intro)
    pdf.ln(2)

    for h, bullets in SECTIONS:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(20, 50, 100)
        pdf.multi_cell(0, 5.5, sanitize(h))
        pdf.ln(0.5)
        pdf.set_font("Helvetica", "", 7.5)
        pdf.set_text_color(30)
        for b in bullets:
            pdf.set_x(pdf.l_margin + 2)
            pdf.multi_cell(0, 4, f"- {sanitize(b)}")
        pdf.ln(1.2)

    pdf.set_font("Helvetica", "I", 6.5)
    pdf.set_text_color(100)
    pdf.multi_cell(0, 3.5, sanitize("Key sources: TradingGroup (2508.17565), MetaAgent (2508.00271), Self-Evolving Agents Survey (2507.21046), TradingAgents reflection loop, Reflexion, HyperAgents, internal tredo architecture (episodes, regret, meta-control, debate, DisciplinedCore). Not financial advice."))
    
    out = os.path.join(OUT_DIR, "self-evolving-agentic-ai-trading-intact-blueprint-2026.pdf")
    pdf.output(out)
    print(f"PDF created: {out}")
    return out

def make_docx():
    doc = Document()
    title = doc.add_heading(sanitize(TITLE), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"Date: {DATE}\n").italic = True
    p.add_run("Research on self-evolving agentic AI and trading systems focused on learning from mistakes, reflection, meta-learning, regret-driven adaptation, and continual self-improvement. Includes diagnosis of current gaps and a complete 'intact system' blueprint.").italic = True

    for h, bullets in SECTIONS:
        doc.add_heading(sanitize(h), level=1)
        for b in bullets:
            doc.add_paragraph(sanitize(b), style="List Bullet")

    doc.add_heading("Key References", level=1)
    refs = [
        "TradingGroup: Multi-Agent Trading System with Self-Reflection and Data-Synthesis (arXiv 2508.17565)",
        "MetaAgent: Toward Self-Evolving Agent via Tool Meta-Learning (arXiv 2508.00271)",
        "A Survey of Self-Evolving Agents (arXiv 2507.21046)",
        "TradingAgents (memory + realized return → reflection injection loop)",
        "Reflexion, HyperAgents/DGM-H, EvoAgentX and experience-driven lifelong learning frameworks",
        "Internal tredo/TREDO: episodes with regret/reflection, DisciplinedCore, debate, trained memory, temporal loops, reflector"
    ]
    for r in refs:
        doc.add_paragraph(sanitize(r), style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run("\nNot financial advice. For design and implementation of safe, self-improving autonomous agentic trading systems.")
    run.italic = True

    out = os.path.join(OUT_DIR, "self-evolving-agentic-ai-trading-intact-blueprint-2026.docx")
    doc.save(out)
    print(f"DOCX created: {out}")
    return out

if __name__ == "__main__":
    pdf = make_pdf()
    docx = make_docx()
    print(f"\nSelf-evolving / intact system reports generated in {OUT_DIR}")
